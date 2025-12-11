import pandas as pd
from docx import Document
from datetime import datetime
import os
from docx.enum.text import WD_BREAK # NECESARIO para el salto de página

# --- 1. FUNCIÓN DE UTILIDAD: Reemplazo de marcadores ---
def reemplazar_marcadores(doc, marcador, valor):
    """Reemplaza un marcador de texto en párrafos y tablas del documento."""
    
    # 1. En Párrafos
    for p in doc.paragraphs:
        if marcador in p.text:
            p.text = p.text.replace(marcador, valor)

    # 2. En Celdas de Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if marcador in cell.text:
                    cell.text = cell.text.replace(marcador, valor)

# --- 2. FUNCIÓN DE ANÁLISIS: Generación de Resúmenes y Valoración ---
def analizar_stock_y_generar_texto(df_inventario):
    """
    Analiza el DataFrame y genera los strings de resumen, valoración y conclusiones.
    """
    
    columnas_analisis = ['Stock Mínimo Deseado', 'Stock Máximo Permitido', 
                         'Nombre de Caja/Producto', 'Cantidad Final (Actual)', 
                         'Unidad de Medida (Ej. Pza, Caja de 12, Kg)', 'Sección/Ubicación Física']
    
    # Validación de columnas para el análisis crítico
    for col in columnas_analisis:
        if col not in df_inventario.columns:
            return {
                'valoracion': "ADVERTENCIA: Faltan datos para análisis.",
                'criticos': f"Advertencia: Análisis incompleto. Falta la columna '{col}' en el Excel.",
                'sobre': "Advertencia: Análisis incompleto.",
                'conclusiones': "No se pudieron generar conclusiones automáticas debido a datos faltantes."
            }
            
    # --- 2.1. Valoración del Inventario ---
    total_productos = len(df_inventario)
    total_unidades_final = df_inventario['Cantidad Final (Actual)'].sum()
    
    if 'Costo Unitario' in df_inventario.columns:
        # Cálculo del valor total (si la columna existe)
        df_inventario['Valor Total'] = df_inventario['Cantidad Final (Actual)'] * df_inventario['Costo Unitario']
        valor_total_inventario = df_inventario['Valor Total'].sum()
        valoracion_texto = (
            f"El inventario actual (Corte {datetime.now().strftime('%Y-%m-%d')}) consta de {total_productos} productos únicos "
            f"con un total de {total_unidades_final:,.2f} unidades. "
            f"El valor total estimado del stock es de ${valor_total_inventario:,.2f} (Calculado con Costo Unitario)."
        )
    else:
        # Mensaje si falta la columna de costo
        valoracion_texto = (
            f"El inventario actual consta de {total_productos} productos únicos, con un total de "
            f"{total_unidades_final:,.2f} unidades. "
            f"NOTA: La valoración monetaria no pudo ser calculada por falta de la columna 'Costo Unitario' en el Excel."
        )

    # --- 2.2. Análisis de Artículos Críticos y Sobre-Inventario (Texto) ---
    df_criticos = df_inventario[df_inventario['Cantidad Final (Actual)'] < df_inventario['Stock Mínimo Deseado']].copy()
    df_sobre = df_inventario[df_inventario['Cantidad Final (Actual)'] > df_inventario['Stock Máximo Permitido']].copy()

    # Generación de texto Críticos (mismo código que antes)
    if df_criticos.empty:
        criticos_texto = "No se detectaron artículos bajo el Stock Mínimo Deseado. ¡Inventario en orden!"
    else:
        criticos_texto = "Se requiere **pedido urgente** de los siguientes artículos:\n\n"
        criticos_lista = []
        for index, row in df_criticos.iterrows():
            diferencia = row['Stock Mínimo Deseado'] - row['Cantidad Final (Actual)']
            criticos_lista.append(
                f"\t- **{row['Nombre de Caja/Producto']}**: Stock actual {row['Cantidad Final (Actual)']} (Faltan {diferencia})."
            )
        criticos_texto += "\n".join(criticos_lista)
        
    # Generación de texto Sobre-Inventario (mismo código que antes)
    if df_sobre.empty:
        sobre_texto = "No se detectó sobre-inventario. La gestión de stock es eficiente."
    else:
        sobre_texto = "Se detectó **exceso de stock** en los siguientes artículos:\n\n"
        sobre_lista = []
        for index, row in df_sobre.iterrows():
            exceso = row['Cantidad Final (Actual)'] - row['Stock Máximo Permitido']
            sobre_lista.append(
                f"\t- **{row['Nombre de Caja/Producto']}**: Stock actual {row['Cantidad Final (Actual)']} (Exceso de {exceso})."
            )
        sobre_texto += "\n".join(sobre_lista)

    # --- 2.3. Generación de Conclusiones Automáticas ---
    num_criticos = len(df_criticos)
    num_sobre = len(df_sobre)
    
    if num_criticos == 0 and num_sobre == 0:
        conclusiones_texto = (
            "CONCLUSIÓN: La gestión de inventario ha sido óptima en este periodo, sin identificar desbalances de stock. "
            "Se sugiere mantener los procedimientos de control y los niveles de Stock Mínimo/Máximo establecidos."
        )
    elif num_criticos > 0 or num_sobre > 0:
        acciones = []
        if num_criticos > 0:
            acciones.append(f"1) Generar una orden de compra urgente para los {num_criticos} artículos críticos.")
        if num_sobre > 0:
            acciones.append(f"2) Reducir los próximos pedidos para los {num_sobre} artículos con sobre-inventario.")
            
        conclusiones_texto = (
            f"CONCLUSIÓN: Se detectó un desbalance en el inventario ({num_criticos} críticos y {num_sobre} con sobre-stock). "
            "Acciones Correctivas Sugeridas: " + " ".join(acciones) + 
            " Se recomienda revisar la planificación de demanda para el próximo mes."
        )
    else:
         conclusiones_texto = "No se pudieron generar conclusiones automáticas."

    return {
        'valoracion': valoracion_texto,
        'criticos': criticos_texto,
        'sobre': sobre_texto,
        'conclusiones': conclusiones_texto
    }

# --- 3. FUNCIÓN DE TABLA: Llenado de la tabla de Inventario ---
def llenar_tabla_inventario(doc, df_inventario):
    """
    Busca la tabla de Inventario Mensual (asumida como la tercera tabla) y la rellena con datos.
    """
    try:
        tabla_inventario = doc.tables[2]
    except IndexError:
        return "Error: No se encontró la tercera tabla (Inventario Mensual) en la plantilla. Verifique que la plantilla tenga 3 tablas."
    
    # Columnas del Excel que coinciden con las del Word (excluyendo 'No.')
    columnas_word = [
        'Nombre de Caja/Producto', 'Código/SKU (si aplica)', 'Sección/Ubicación Física', 
        'Unidad de Medida (Ej. Pza, Caja de 12, Kg)', 'Cantidad Inicial (al inicio del mes)', 
        'Entradas (Total del Mes)', 'Salidas (Total del Mes)', 'Cantidad Final (Actual)', 
        'Diferencia (Actual vs. Teórico)', 'Causa de la Diferencia (Si existe)', 
        'Observaciones y Estatus'
    ]

    # 1. Limpiar las filas de ejemplo de la plantilla (todas menos la de encabezado)
    while len(tabla_inventario.rows) > 1:
        row_to_delete = tabla_inventario.rows[-1]
        row_to_delete._element.getparent().remove(row_to_delete._element)
    
    # 2. Rellenar las filas con los datos del DataFrame
    for index, row in df_inventario.iterrows():
        cells = tabla_inventario.add_row().cells
        cells[0].text = str(index + 1)
        
        # Llenar las columnas de datos (índices 1 a 11)
        for i, col_name in enumerate(columnas_word):
            if col_name in row.index:
                try:
                    # Formateo de números grandes para mejor lectura
                    valor = row[col_name]
                    if isinstance(valor, (int, float)) and abs(valor) > 1000:
                        valor = f"{valor:,.2f}"
                    cells[i + 1].text = str(valor)
                except:
                    cells[i + 1].text = ""
            else:
                cells[i + 1].text = '' 

    return None

# Dentro de reporte_logica.py

# ... (Funciones auxiliares reemplazar_marcadores, analizar_stock_y_generar_texto, llenar_tabla_inventario son las MISMAS) ...

# --- 4. FUNCIÓN DE ORQUESTACIÓN: Ejecución completa (Modificada) ---
def generar_reporte_completo(ruta_excel, ruta_plantilla, datos_gui):
    """
    Orquesta la lectura de datos, el análisis, el llenado de la plantilla
    y guarda el reporte final, usando los datos proporcionados por la GUI.
    """
    try:
        # 1. Lectura de Datos del Excel
        df_inventario = pd.read_excel(ruta_excel, sheet_name=0) 
        
        # 2. Cargar la Plantilla de Word
        if not os.path.exists(ruta_plantilla):
            return False, f"Error de Plantilla: No se encontró el archivo en la ruta: {ruta_plantilla}"
            
        doc = Document(ruta_plantilla)
        
        # 3. Sustitución de Marcadores de Fecha y Responsabilidades usando datos_gui
        
        # Fechas
        reemplazar_marcadores(doc, '<<MES>>', datos_gui['mes'])
        reemplazar_marcadores(doc, '<<FECHA_INICIO>>', datos_gui['fecha_inicio'])
        reemplazar_marcadores(doc, '<<FECHA_FIN>>', datos_gui['fecha_fin'])
        
        # Responsabilidades
        reemplazar_marcadores(doc, '<<NOMBRE_DIR>>', datos_gui['dir_nombre'])
        reemplazar_marcadores(doc, '<<NOMBRE_RESP>>', datos_gui['resp_nombre'])
        reemplazar_marcadores(doc, '<<NOMBRE_VERIF>>', datos_gui['verif_nombre'])
        
        # 4. Análisis Dinámico y Llenado de Resumen (Sin cambios, usa df_inventario)
        resumenes = analizar_stock_y_generar_texto(df_inventario)
        
        reemplazar_marcadores(doc, '<<VALORACION_INVENTARIO>>', resumenes['valoracion'])
        reemplazar_marcadores(doc, '<<CRITICOS_RESUMEN>>', resumenes['criticos'])
        reemplazar_marcadores(doc, '<<SOBRE_INVENTARIO_RESUMEN>>', resumenes['sobre'])
        reemplazar_marcadores(doc, '<<CONCLUSIONES>>', resumenes['conclusiones'])
        
        # 5. Insertar Salto de Página antes de la tabla de Inventario Mensual
        # (El mismo código para el salto de página que ya implementamos)
        from docx.enum.text import WD_BREAK # Asegurar que esté importado aquí si no lo está al inicio
        
        salto_insertado = False
        for p in doc.paragraphs:
            if 'El registro debe ser llenado por' in p.text:
                p.runs[0].add_break(WD_BREAK.PAGE)
                salto_insertado = True
                break
        
        if not salto_insertado:
             # Opción de respaldo: si no encuentra la frase anterior, busca el título "Inventario Mensual"
            for p in doc.paragraphs:
                if 'Inventario Mensual' in p.text:
                     p.insert_paragraph_before('').runs[0].add_break(WD_BREAK.PAGE)
                     break

        # 6. Llenar la Tabla de Inventario Mensual
        error_tabla = llenar_tabla_inventario(doc, df_inventario)
        if error_tabla:
            return False, error_tabla

        # 7. Guardar el Documento Final
        ahora = datetime.now()
        nombre_final = f"Reporte_Inventario_{ahora.strftime('%Y%m%d_%H%M%S')}.docx"
        ruta_salida = os.path.join(os.path.dirname(ruta_excel), nombre_final)
        doc.save(ruta_salida)
        
        return True, ruta_salida

    except FileNotFoundError:
        return False, "Error: Archivo de Excel no encontrado. Verifique la ruta y el nombre."
    except Exception as e:
        return False, f"Error desconocido durante el procesamiento: {type(e).__name__} - {e}"